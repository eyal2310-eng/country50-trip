"""
עזר להפקת מסמכי Word בעברית עם RTL מלא.

שימוש:
    from rtl_docx import new_doc, para, heading, table, save
    doc = new_doc()
    heading(doc, "חוות דעת משפטית", level=1)
    para(doc, "בשם מרשי, הריני לפנות אליך כדלקמן:")
    para(doc, "1. העובדה הראשונה.")
    save(doc, "חוות-דעת-כהן-2026-09-03.docx")

דורש: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FONT = "David"
SIZE = 12
HEADING_SIZES = {1: 16, 2: 14, 3: 12}


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def _rtl_rpr(rPr, font, size):
    rPr.append(_el("w:rtl", **{"w:val": "1"}))
    rPr.get_or_add_rFonts().set(qn("w:cs"), font)
    rPr.append(_el("w:szCs", **{"w:val": str(int(size * 2))}))
    rPr.append(_el("w:lang", **{"w:bidi": "he-IL"}))


def new_doc(font=FONT, size=SIZE, margin_cm=2.5):
    """מסמך חדש שכולו RTL כברירת מחדל."""
    doc = Document()
    for section in doc.sections:
        section._sectPr.append(_el("w:bidi"))
        section.left_margin = section.right_margin = Cm(margin_cm)
        section.top_margin = section.bottom_margin = Cm(margin_cm)
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(size)
    _rtl_rpr(style.element.get_or_add_rPr(), font, size)
    pPr = style.element.get_or_add_pPr()
    pPr.append(_el("w:bidi", **{"w:val": "1"}))
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


def _fix_para(p, align, font, size, bold):
    p._p.get_or_add_pPr().append(_el("w:bidi", **{"w:val": "1"}))
    p.alignment = align
    for run in p.runs:
        run.font.name = font
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        _rtl_rpr(run._r.get_or_add_rPr(), font, size)
    return p


def para(doc, text="", bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         font=FONT, size=SIZE):
    return _fix_para(doc.add_paragraph(text), align, font, size, bold)


def heading(doc, text, level=1, font=FONT):
    size = HEADING_SIZES.get(level, SIZE)
    p = doc.add_heading(text, level=level)
    return _fix_para(p, WD_ALIGN_PARAGRAPH.RIGHT, font, size, True)


def bullet(doc, text, numbered=False, font=FONT, size=SIZE):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(text, style=style)
    return _fix_para(p, WD_ALIGN_PARAGRAPH.JUSTIFY, font, size, None)


def table(doc, rows, header=True, font=FONT, size=SIZE):
    """rows = רשימת רשימות. השורה הראשונה היא כותרת אם header=True."""
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t._tbl.tblPr.append(_el("w:bidiVisual"))
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = str(val)
            for p in cell.paragraphs:
                _fix_para(p, WD_ALIGN_PARAGRAPH.RIGHT, font, size,
                          True if (header and i == 0) else None)
    return t


def footer(doc, text, font=FONT, size=9):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.text = text
        _fix_para(p, WD_ALIGN_PARAGRAPH.CENTER, font, size, None)


def save(doc, path):
    doc.save(path)
    return path
